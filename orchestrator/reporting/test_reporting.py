"""Runnable self-test for the reporting engine.

Plain asserts + ``__main__``. Writes everything into a temp dir and cleans up.

Run:  python orchestrator/reporting/test_reporting.py
"""

from __future__ import annotations

import os
import sys
import tempfile

# Allow running as a plain script (python orchestrator/reporting/test_reporting.py)
# as well as a module, by ensuring the package's parent is importable.
_HERE = os.path.dirname(os.path.abspath(__file__))
_ORCH = os.path.dirname(_HERE)
if _ORCH not in sys.path:
    sys.path.insert(0, _ORCH)

from openpyxl import load_workbook
from docx import Document
import pikepdf

from reporting.export import (
    build_xlsx, build_docx, build_pdf, build_reports,
    _REGISTER_HEADERS,
)


# A small fixture mirroring the seed.py finding shape.
FINDINGS = [
    {
        "id": "VLN-2087", "title": "BOLA on /v1/claims/{id}",
        "severity": "critical", "status": "triaged", "asset": "Claims Processing API",
        "framework": "OWASP API", "catCode": "API1:2023", "cvss": 9.3,
        "deadline": "2026-06-04", "daysLeft": 3, "isClosed": False, "owner": "R. Iyer",
    },
    {
        "id": "VLN-2074", "title": "Default admin credentials",
        "severity": "critical", "status": "open", "asset": "Document Management Server",
        "framework": "CIS", "catCode": "CIS-5.2", "cvss": 9.8,
        "deadline": "2026-05-25", "daysLeft": -7, "isClosed": False, "owner": "Unassigned",
    },
    {
        "id": "VLN-2052", "title": "Missing rate limiting on OTP",
        "severity": "high", "status": "triaged", "asset": "Agent Mobile App (API)",
        "framework": "OWASP API", "catCode": "API4:2023", "cvss": 7.4,
        "deadline": "2026-05-28", "daysLeft": -4, "isClosed": False, "owner": "S. Khan",
    },
    {
        "id": "VLN-2019", "title": "Session cookie missing flags",
        "severity": "medium", "status": "open", "asset": "Internal HR Portal",
        "framework": "OWASP Web", "catCode": "A05:2021", "cvss": 5.1,
        "deadline": "2026-07-02", "daysLeft": 31, "isClosed": False, "owner": "Unassigned",
    },
    {
        "id": "VLN-1990", "title": "Banner discloses server version",
        "severity": "info", "status": "open", "asset": "Branch VPN Gateway",
        "framework": "CIS", "catCode": "CIS-4.8", "cvss": 2.0,
        "deadline": None, "daysLeft": None, "isClosed": False, "owner": "Unassigned",
    },
    {
        "id": "VLN-1965", "title": "Open redirect on login return URL",
        "severity": "low", "status": "closed", "asset": "Policyholder Portal",
        "framework": "OWASP Web", "catCode": "A01:2021", "cvss": 3.4,
        "deadline": "2026-04-01", "daysLeft": -61, "isClosed": True, "owner": "A. Mehta",
    },
]

OPEN_PW = "open-me-1234"
OWNER_PW = "owner-secret-9876"
META = {"title": "Vantage Findings — Test", "scan_id": "SCAN-0098"}


def _nonempty(path: str) -> None:
    assert os.path.exists(path), f"missing: {path}"
    assert os.path.getsize(path) > 0, f"empty: {path}"


def test_xlsx(path: str) -> None:
    build_xlsx(FINDINGS, path, meta=META)
    _nonempty(path)

    wb = load_workbook(path)
    assert "Findings" in wb.sheetnames, wb.sheetnames
    assert "SLA summary" in wb.sheetnames, wb.sheetnames

    ws = wb["Findings"]
    # header row matches
    header = [c.value for c in ws[1]]
    assert header == _REGISTER_HEADERS, header
    assert "Severity" in header
    # register row count = 1 header + N findings
    assert ws.max_row == len(FINDINGS) + 1, ws.max_row

    # SLA summary: critical count == 2 and an overdue line is present.
    sla = wb["SLA summary"]
    rows = {r[0].value: r[1].value for r in sla.iter_rows() if r[0].value}
    assert rows.get("critical") == 2, rows
    assert "Overdue (SLA breached)" in rows, rows
    wb.close()
    print("  [ok] xlsx: Findings + SLA summary sheets, header + row count verified")


def test_docx(path: str) -> None:
    build_docx(FINDINGS, path, meta=META)
    _nonempty(path)

    doc = Document(path)
    texts = [p.text for p in doc.paragraphs]
    assert any(META["title"] in t for t in texts), "title heading missing"
    assert any(t.startswith("Generated:") for t in texts), "generated date missing"

    assert len(doc.tables) >= 1, "no findings table"
    table = doc.tables[0]
    # header row + one row per finding
    assert len(table.rows) == len(FINDINGS) + 1, len(table.rows)
    hdr = [c.text for c in table.rows[0].cells]
    assert hdr == _REGISTER_HEADERS, hdr
    print("  [ok] docx: title, generated-date, findings table present")


def test_pdf(path: str) -> None:
    build_pdf(FINDINGS, path, OPEN_PW, OWNER_PW, meta=META)
    _nonempty(path)

    # 1. No password -> PasswordError.
    raised = False
    try:
        pikepdf.open(path)
    except pikepdf.PasswordError:
        raised = True
    assert raised, "PDF opened with NO password (not encrypted!)"

    # 2. Open with the user/open password -> succeeds, and copy/modify blocked.
    with pikepdf.open(path, password=OPEN_PW) as pdf:
        assert pdf.is_encrypted, "pdf reports not encrypted"
        allow = pdf.allow
        assert allow.extract is False, "copy/extract NOT disabled"
        assert allow.modify_other is False, "modify NOT disabled"
        assert allow.modify_assembly is False, "assembly/modify NOT disabled"
        assert allow.print_highres is False, "print NOT disabled"
        assert pdf.encryption.R == 6, f"expected AES-256 R6, got R={pdf.encryption.R}"

    # 3. Owner password also opens the document.
    with pikepdf.open(path, password=OWNER_PW) as pdf:
        assert pdf.is_encrypted
    print("  [ok] pdf: no-pw -> PasswordError; open-pw works; copy/modify/print disabled (AES-256/R6)")


def test_bundle(outdir: str) -> None:
    bundle = build_reports(FINDINGS, outdir, OPEN_PW, OWNER_PW, meta=META)
    assert set(bundle) == {"xlsx", "docx", "pdf"}, bundle
    for p in bundle.values():
        _nonempty(p)
    # the bundle PDF is also dual-password protected.
    raised = False
    try:
        pikepdf.open(bundle["pdf"])
    except pikepdf.PasswordError:
        raised = True
    assert raised, "bundle PDF not encrypted"
    print(f"  [ok] build_reports wrote: {list(bundle)}")


def main() -> int:
    with tempfile.TemporaryDirectory(prefix="vantage-report-test-") as td:
        print("temp dir:", td)
        test_xlsx(os.path.join(td, "f.xlsx"))
        test_docx(os.path.join(td, "f.docx"))
        test_pdf(os.path.join(td, "f.pdf"))
        test_bundle(os.path.join(td, "bundle"))
    # temp dir (and all files) cleaned up on context exit.
    print("ALL REPORTING TESTS PASSED")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
