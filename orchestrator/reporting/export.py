"""Report generators for Vantage.

Three artifacts from a list of plain finding dicts (shape per
``orchestrator/api/seed.py``: id, title, severity, status, asset, cvss,
framework/catCode, deadline, daysLeft, owner, ...):

  * build_xlsx — openpyxl: "Findings" register + "SLA summary" sheets.
  * build_docx — python-docx: SOP-style report (title, date, summary, table).
  * build_pdf  — reportlab content, then pikepdf AES-256 dual-password encrypt.

``build_reports`` writes all three and returns their paths.

Design notes
------------
PDF protection. reportlab is used for the *content* because it is pure-Python
and needs no system libraries (unlike weasyprint, which pulls in pango/cairo).
The encryption is applied by pikepdf as a second pass: we render an unencrypted
PDF in memory, then save it through ``pikepdf.Encryption`` with a distinct
*user* (open) password and *owner* password, AES-256 (R=6, 256-bit), with a
``pikepdf.Permissions`` object that disables copy/extract, modification and
printing. Anyone opening the file must supply the open password; even with it,
the owner-restricted operations (copy/modify/print) stay disabled until the
owner password is supplied.
"""

from __future__ import annotations

import io
import os
from collections import Counter
from datetime import datetime, timezone
from typing import Iterable, Optional

# ----- Excel -----
from openpyxl import Workbook, load_workbook  # noqa: F401 (load used by tests)
from openpyxl.styles import Font, PatternFill

# ----- Word -----
from docx import Document
from docx.shared import Pt

# ----- PDF content + encryption -----
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
)
import pikepdf


SEVERITY_ORDER = ["critical", "high", "medium", "low", "info"]

# Register columns (header row + the field each cell reads).
_REGISTER_HEADERS = [
    "ID", "Severity", "Status", "Asset", "Category",
    "CVSS", "Deadline", "Days Left", "Owner",
]


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------
def _g(f: dict, *keys, default=""):
    """First present, non-None value among ``keys`` (else ``default``)."""
    for k in keys:
        v = f.get(k)
        if v is not None:
            return v
    return default


def _category(f: dict) -> str:
    """Human category label: 'FRAMEWORK CODE' (e.g. 'OWASP API API1:2023')."""
    fw = _g(f, "framework", default="")
    code = _g(f, "catCode", "code", default="")
    return " ".join(p for p in (str(fw), str(code)) if p).strip()


def _register_row(f: dict) -> list:
    return [
        _g(f, "id"),
        _g(f, "severity"),
        _g(f, "status"),
        _g(f, "asset"),
        _category(f),
        _g(f, "cvss", default=""),
        _g(f, "deadline", default="—"),
        _g(f, "daysLeft", default="—"),
        _g(f, "owner", default="Unassigned"),
    ]


def _severity_counts(findings: Iterable[dict]) -> "Counter[str]":
    c: "Counter[str]" = Counter()
    for f in findings:
        c[str(_g(f, "severity", default="unknown")).lower()] += 1
    return c


def _overdue_count(findings: Iterable[dict]) -> int:
    """Open findings whose SLA window has elapsed (daysLeft < 0)."""
    n = 0
    for f in findings:
        if _g(f, "isClosed", default=False):
            continue
        dl = f.get("daysLeft")
        if isinstance(dl, (int, float)) and dl < 0:
            n += 1
    return n


def _now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


def _title(meta: Optional[dict]) -> str:
    if meta and meta.get("title"):
        return str(meta["title"])
    return "Vantage Vulnerability Findings Report"


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------
def build_xlsx(findings: list, path: str, meta: Optional[dict] = None) -> str:
    """Write the findings register + SLA summary workbook to ``path``."""
    wb = Workbook()

    # --- Sheet 1: Findings register ---
    ws = wb.active
    ws.title = "Findings"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="1F3A5F")
    ws.append(_REGISTER_HEADERS)
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill

    for f in findings:
        ws.append(_register_row(f))

    # Reasonable column widths.
    widths = [12, 10, 12, 28, 30, 7, 13, 11, 16]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = w
    ws.freeze_panes = "A2"

    # --- Sheet 2: SLA summary ---
    sla = wb.create_sheet("SLA summary")
    sla.append(["Severity", "Count"])
    for cell in sla[1]:
        cell.font = header_font
        cell.fill = header_fill

    counts = _severity_counts(findings)
    for sev in SEVERITY_ORDER:
        sla.append([sev, counts.get(sev, 0)])
    # any severities outside the known set
    for sev in sorted(set(counts) - set(SEVERITY_ORDER)):
        sla.append([sev, counts[sev]])

    sla.append([])
    sla.append(["Total findings", sum(counts.values())])
    sla.append(["Overdue (SLA breached)", _overdue_count(findings)])
    sla.column_dimensions["A"].width = 26
    sla.column_dimensions["B"].width = 10

    wb.save(path)
    return path


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------
def build_docx(findings: list, path: str, meta: Optional[dict] = None) -> str:
    """Write a SOP-style Word report to ``path``."""
    doc = Document()

    doc.add_heading(_title(meta), level=0)

    p = doc.add_paragraph()
    run = p.add_run(f"Generated: {_now_iso()}")
    run.italic = True
    run.font.size = Pt(9)
    if meta and meta.get("scan_id"):
        doc.add_paragraph(f"Scan: {meta['scan_id']}")

    # Summary — counts by severity.
    doc.add_heading("Summary", level=1)
    counts = _severity_counts(findings)
    summary = ", ".join(
        f"{sev.capitalize()}: {counts.get(sev, 0)}" for sev in SEVERITY_ORDER
    )
    doc.add_paragraph(
        f"Total findings: {sum(counts.values())}. "
        f"Overdue (SLA breached): {_overdue_count(findings)}."
    )
    doc.add_paragraph(f"By severity — {summary}.")

    # Findings table.
    doc.add_heading("Findings", level=1)
    table = doc.add_table(rows=1, cols=len(_REGISTER_HEADERS))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(_REGISTER_HEADERS):
        hdr[i].text = h
        for par in hdr[i].paragraphs:
            for r in par.runs:
                r.bold = True

    for f in findings:
        cells = table.add_row().cells
        for i, val in enumerate(_register_row(f)):
            cells[i].text = "" if val is None else str(val)

    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# LaTeX (.tex) — a compilable source report
# ---------------------------------------------------------------------------
# Finding data (titles, asset names, owners) is untrusted-ish text that can
# contain LaTeX-special characters, so EVERY cell is escaped before it reaches
# the document. Order matters: backslash is handled per-character (its
# replacement is emitted verbatim, never re-escaped).
_TEX_SPECIAL = {
    "\\": r"\textbackslash{}", "&": r"\&", "%": r"\%", "$": r"\$",
    "#": r"\#", "_": r"\_", "{": r"\{", "}": r"\}",
    "~": r"\textasciitilde{}", "^": r"\textasciicircum{}",
}


def _tex(value) -> str:
    """Escape a value for safe inclusion in LaTeX text."""
    return "".join(_TEX_SPECIAL.get(ch, ch) for ch in str(value))


def build_tex(findings: list, path: str, meta: Optional[dict] = None) -> str:
    """Write a compilable LaTeX (.tex) findings report to ``path``.

    Self-contained article (geometry + longtable + booktabs); compiles with
    pdflatex. Same register + severity summary as the other formats, with every
    data cell LaTeX-escaped.
    """
    counts = _severity_counts(findings)
    title = _tex(_title(meta))

    lines: list[str] = []
    ap = lines.append
    ap(r"\documentclass[10pt,a4paper]{article}")
    ap(r"\usepackage[T1]{fontenc}")
    ap(r"\usepackage[utf8]{inputenc}")
    ap(r"\usepackage[margin=1.8cm,landscape]{geometry}")
    ap(r"\usepackage{longtable}")
    ap(r"\usepackage{booktabs}")
    ap(r"\usepackage{array}")
    ap(r"\usepackage{xcolor}")
    ap(r"\newcolumntype{L}[1]{>{\raggedright\arraybackslash}p{#1}}")
    ap(r"\title{" + title + r"}")
    ap(r"\date{" + _tex(_now_iso()) + r"}")
    ap(r"\begin{document}")
    ap(r"\maketitle")
    if meta and meta.get("scan_id"):
        ap(r"\noindent\textbf{Scan:} " + _tex(meta["scan_id"]) + r"\\[4pt]")
    # Summary
    total = sum(counts.values())
    ap(r"\noindent\textbf{Total findings:} " + str(total)
       + r" \quad \textbf{Overdue (SLA breached):} " + str(_overdue_count(findings)) + r"\\[2pt]")
    sev_summary = ", ".join(
        f"{sev.capitalize()}: {counts.get(sev, 0)}" for sev in SEVERITY_ORDER
    )
    ap(r"\noindent\textbf{By severity:} " + _tex(sev_summary) + r"\par\vspace{8pt}")
    # Register table
    ap(r"{\scriptsize")
    ap(r"\begin{longtable}{l l l L{4.2cm} L{4.2cm} r l r L{2.6cm}}")
    ap(r"\toprule")
    ap(" & ".join(r"\textbf{" + _tex(h) + r"}" for h in _REGISTER_HEADERS) + r" \\")
    ap(r"\midrule")
    ap(r"\endhead")
    for f in findings:
        cells = [_tex(v) for v in _register_row(f)]
        ap(" & ".join(cells) + r" \\")
    ap(r"\bottomrule")
    ap(r"\end{longtable}")
    ap(r"}")
    ap(r"\end{document}")

    text = "\n".join(lines) + "\n"
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(text)
    return path


# ---------------------------------------------------------------------------
# PDF — reportlab content, pikepdf dual-password encryption
# ---------------------------------------------------------------------------
def _render_pdf_bytes(findings: list, meta: Optional[dict]) -> bytes:
    """Render the (unencrypted) PDF content with reportlab; return bytes."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=15 * mm, rightMargin=15 * mm,
        topMargin=15 * mm, bottomMargin=15 * mm,
        title=_title(meta),
    )
    styles = getSampleStyleSheet()
    cell_style = ParagraphStyle(
        "cell", parent=styles["BodyText"], fontSize=7.5, leading=9,
    )
    head_style = ParagraphStyle(
        "cellhead", parent=cell_style, textColor=colors.white, fontSize=8,
    )

    story = []
    story.append(Paragraph(_title(meta), styles["Title"]))
    story.append(Paragraph(f"Generated: {_now_iso()}", styles["Normal"]))
    if meta and meta.get("scan_id"):
        story.append(Paragraph(f"Scan: {meta['scan_id']}", styles["Normal"]))
    story.append(Spacer(1, 6 * mm))

    counts = _severity_counts(findings)
    summary = ", ".join(
        f"{sev.capitalize()}: {counts.get(sev, 0)}" for sev in SEVERITY_ORDER
    )
    story.append(Paragraph(
        f"<b>Total findings:</b> {sum(counts.values())} &nbsp;&nbsp; "
        f"<b>Overdue (SLA breached):</b> {_overdue_count(findings)}",
        styles["Normal"],
    ))
    story.append(Paragraph(f"<b>By severity:</b> {summary}", styles["Normal"]))
    story.append(Spacer(1, 6 * mm))

    # Findings table.
    data = [[Paragraph(h, head_style) for h in _REGISTER_HEADERS]]
    for f in findings:
        data.append([
            Paragraph("" if v is None else str(v), cell_style)
            for v in _register_row(f)
        ])

    col_widths = [
        18 * mm, 15 * mm, 18 * mm, 30 * mm, 35 * mm,
        10 * mm, 16 * mm, 13 * mm, 22 * mm,
    ]
    table = Table(data, colWidths=col_widths, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F3A5F")),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B0B7C3")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#F2F5F9")]),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))
    story.append(table)

    doc.build(story)
    return buf.getvalue()


def build_pdf(
    findings: list,
    path: str,
    open_password: str,
    owner_password: str,
    meta: Optional[dict] = None,
) -> str:
    """Render the report and write it AES-256 encrypted with two passwords.

    ``open_password``  -> the *user* password (required to open the file).
    ``owner_password`` -> the *owner* password (required to lift the
    restrictions below). They must differ for the owner password to mean
    anything.

    Permissions disable copy/extract, all modification, and printing; reading
    (once the open password is supplied) remains allowed.
    """
    if not open_password or not owner_password:
        raise ValueError("Both open_password and owner_password are required.")
    if open_password == owner_password:
        raise ValueError(
            "open_password and owner_password must differ "
            "(otherwise opening the file also grants owner rights)."
        )

    pdf_bytes = _render_pdf_bytes(findings, meta)

    permissions = pikepdf.Permissions(
        extract=False,            # block copy / text & graphics extraction
        modify_annotation=False,  # no annotation edits
        modify_assembly=False,    # no page insert/delete/rotate
        modify_form=False,        # no form filling
        modify_other=False,       # no other modification
        print_lowres=False,       # block printing (low res)
        print_highres=False,      # block printing (high res)
    )
    encryption = pikepdf.Encryption(
        user=open_password,       # password to OPEN the document
        owner=owner_password,     # password to change permissions
        R=6,                      # AES-256 (PDF 2.0 / R6)
        allow=permissions,
    )

    with pikepdf.open(io.BytesIO(pdf_bytes)) as pdf:
        pdf.save(path, encryption=encryption)
    return path


# ---------------------------------------------------------------------------
# Bundle
# ---------------------------------------------------------------------------
def build_reports(
    findings: list,
    outdir: str,
    open_password: str,
    owner_password: str,
    meta: Optional[dict] = None,
) -> dict:
    """Write xlsx + docx + dual-password pdf into ``outdir``.

    Returns ``{"xlsx": ..., "docx": ..., "pdf": ...}`` absolute paths.
    """
    os.makedirs(outdir, exist_ok=True)
    base = "vantage-report"
    if meta and meta.get("scan_id"):
        base = f"vantage-{meta['scan_id']}"

    xlsx_path = os.path.abspath(os.path.join(outdir, base + ".xlsx"))
    docx_path = os.path.abspath(os.path.join(outdir, base + ".docx"))
    pdf_path = os.path.abspath(os.path.join(outdir, base + ".pdf"))

    build_xlsx(findings, xlsx_path, meta=meta)
    build_docx(findings, docx_path, meta=meta)
    build_pdf(findings, pdf_path, open_password, owner_password, meta=meta)

    return {"xlsx": xlsx_path, "docx": docx_path, "pdf": pdf_path}
